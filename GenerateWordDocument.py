import re
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    A helper function that places a hyperlink within a paragraph.
    Based on known recipes for adding hyperlinks using python-docx.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color is not None:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)

    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_formatted_text(paragraph, text):
    """
    Scans the given text for inline markdown elements and adds runs to the
    paragraph with the proper formatting.

    Currently supported:
      - **bold text**
      - [link text](url)
    """
    # Regex to match **bold text** or markdown links [text](url)
    pattern = re.compile(
        r"(\*\*(?P<bold>.+?)\*\*|\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            plain_text = text[pos:match.start()]
            paragraph.add_run(plain_text)
        if match.group('bold'):
            bold_text = match.group('bold')
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif match.group('link_text'):
            link_text = match.group('link_text')
            link_url = match.group('link_url')
            add_hyperlink(paragraph, link_url, link_text)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def generate_word_doc_from_markdown(mark_down_text: str) -> Document:
    """
    Converts markdown text into a Word Document object, preserving headers,
    paragraphs, bullet lists, links, and bold formatting.
    """
    doc = Document()
    lines = mark_down_text.splitlines()
    paragraph_lines = []

    def flush_paragraph_lines():
        """Adds any accumulated paragraph text as a new paragraph with inline formatting."""
        if paragraph_lines:
            paragraph_text = " ".join(paragraph_lines)
            p = doc.add_paragraph()
            add_formatted_text(p, paragraph_text)
            paragraph_lines.clear()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_paragraph_lines()
            continue

        # Process markdown headers.
        if stripped.startswith("#"):
            flush_paragraph_lines()
            header_level = 0
            while header_level < len(stripped) and stripped[header_level] == "#":
                header_level += 1
            header_level = min(header_level, 6)
            header_text = stripped[header_level:].strip()
            heading_paragraph = doc.add_heading("", level=header_level)
            add_formatted_text(heading_paragraph, header_text)
        # Process unordered list items: only if the marker is immediately followed by a space.
        elif stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ "):
            flush_paragraph_lines()
            # Remove the marker (first 2 characters: marker and space)
            list_text = stripped[2:].strip()
            list_paragraph = doc.add_paragraph("", style='List Bullet')
            add_formatted_text(list_paragraph, list_text)
        else:
            paragraph_lines.append(stripped)

    flush_paragraph_lines()
    return doc


# Example usage:
if __name__ == '__main__':
    # Example markdown text
    markdown_input = """
    # Curated Articles for Customer Navigator Letter

## [Streamline grant proposal reviews using Amazon Bedrock](https://aws.amazon.com/blogs/machine-learning/streamline-grant-proposal-reviews-using-amazon-bedrock/)
**Category:** Mass Data Summarization Use Cases  
**Takeaway:** **Amazon Bedrock reduces grant proposal review time by 90%:**  
The AWS Machine Learning Blog details how Amazon Bedrock, a generative AI tool, has enabled government and non-profit organizations to streamline grant proposal reviews, reducing processing time from 21 days to 2 days. The prototype application, developed using Streamlit, Amazon Bedrock, and Amazon DynamoDB, allows for dynamic evaluation of proposals based on user-defined personas and rubrics. This innovation highlights the potential for AI-driven efficiencies in administrative processes, urging AI leaders to consider similar applications for reducing manual workloads and enhancing operational efficiency.

**Rationale:**  
- **Emerging trends:** Utilization of generative AI for administrative efficiency.
- **Potential risks:** Scalability and security considerations for production-grade applications.
- **Applications:** Streamlined processes for grant proposal reviews.
- **Key statistics or findings:** 90% reduction in processing time.
- **Actionable recommendations:** Explore AI-driven solutions for other time-intensive administrative tasks.

## [DeepSeek-R1 models now available on AWS](https://aws.amazon.com/blogs/aws/deepseek-r1-models-now-available-on-aws/)
**Category:** Competitors  
**Takeaway:** **AWS expands AI offerings with DeepSeek-R1 models:**  
AWS has introduced DeepSeek-R1 models on its platform, providing cost-effective AI solutions with various deployment options through Amazon Bedrock and Amazon SageMaker. These models emphasize cost efficiency, security, and flexibility, allowing enterprises to tailor AI deployments to their specific needs. AI leaders should assess these models for potential integration into their AI strategies, particularly for cost-sensitive projects.

**Rationale:**  
- **Emerging trends:** Cost-effective AI model deployment.
- **Potential risks:** Ensuring security and compliance in AI deployments.
- **Applications:** Customizable AI solutions for diverse enterprise needs.
- **Key statistics or findings:** Emphasis on cost efficiency and security.
- **Actionable recommendations:** Evaluate DeepSeek-R1 models for projects requiring scalable and secure AI solutions.

## [Researchers develop even cheaper alternative to DeepSeek's R1-Zero model](https://futurism.com/researchers-deepseek-even-cheaper)
**Category:** Small Language Models  
**Takeaway:** **TinyZero challenges AI cost paradigms with $30 model:**  
Jiayi Pan and his team at UC Berkeley have developed TinyZero, a cost-effective alternative to DeepSeek's R1-Zero model, demonstrating AI capabilities at a fraction of the cost. This development highlights the potential for open-source and budget-friendly AI solutions, challenging the need for large financial investments in AI infrastructure. AI leaders should explore these low-cost models for applications where budget constraints are a priority.

**Rationale:**  
- **Emerging trends:** Budget-friendly AI model development.
- **Potential risks:** Performance limitations compared to more expensive models.
- **Applications:** Cost-effective AI solutions for constrained environments.
- **Key statistics or findings:** TinyZero developed for just $30.
- **Actionable recommendations:** Consider integrating low-cost models for non-critical applications to optimize budgets.

## [Authors v. Anthropic: Judge gets tutorial on AI training in copyright battle](https://www.dailyjournal.com/articles/383177-authors-v-anthropic-judge-gets-tutorial-on-ai-training-in-copyright-battle)
**Category:** Explainability & Contestability  
**Takeaway:** **AI copyright infringement raises transparency issues:**  
The case of Bartz v. Anthropic PBC highlights the legal challenges surrounding AI training data, with claims of copyright infringement due to the use of plaintiffs' works in training the Claude language model. This case underscores the importance of transparency and explainability in AI development to mitigate legal risks. AI leaders must prioritize ethical AI practices and ensure compliance with intellectual property laws to avoid similar disputes.

**Rationale:**  
- **Emerging trends:** Legal scrutiny of AI training data.
- **Potential risks:** Intellectual property infringement and legal liabilities.
- **Applications:** Transparent AI development practices.
- **Key statistics or findings:** Ongoing legal challenges for AI companies.
- **Actionable recommendations:** Implement robust compliance checks and transparent data usage policies in AI projects.

## [Artificial intelligence transforming lives of hearing impaired and hard of hearing people](https://www.telefonica.com/en/communication-room/blog/artificial-intelligence-transforming-lives-hearing-impaired-hard-hearing-people/)
**Category:** Conversational AI Assistants Use Cases  
**Takeaway:** **AI enhances accessibility for hearing impaired individuals:**  
AI technologies, such as automatic captioning and AI-enabled hearing aids, are significantly improving communication for the hearing impaired, promoting inclusion and independence. Despite challenges like connectivity and accuracy, these advancements demonstrate AI's potential to transform accessibility. AI leaders should consider incorporating similar technologies to enhance inclusivity and accessibility in their products and services.

**Rationale:**  
- **Emerging trends:** AI-driven accessibility solutions.
- **Potential risks:** Connectivity and accuracy limitations.
- **Applications:** Enhanced communication tools for the hearing impaired.
- **Key statistics or findings:** AI promotes inclusion and independence.
- **Actionable recommendations:** Invest in developing AI solutions that address accessibility challenges for diverse user groups.
"""
    # Generate the document from the markdown text.
    document = generate_word_doc_from_markdown(markdown_input)
    # Save the document to a file.
    document.save("output.docx")
    print("Word document 'output.docx' created successfully.")
